"""
Document Generator - Generates Word, PDF, and includes UML diagrams

Supports:
- Word (.docx) generation with proper formatting
- PDF generation
- UML diagram generation (PlantUML, Mermaid)
- Image embedding
- Table of Contents
- Professional styling
"""

import os
import subprocess
import tempfile
from pathlib import Path
from typing import Optional, List, Dict, Any
from dataclasses import dataclass
from datetime import datetime
import base64
import io

# Check for optional dependencies
DOCX_AVAILABLE = False
PDF_AVAILABLE = False
PILLOW_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    pass

try:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4, letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch, cm
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image, PageBreak
    from reportlab.platypus.tableofcontents import TableOfContents
    PDF_AVAILABLE = True
except ImportError:
    pass

try:
    from PIL import Image as PILImage
    PILLOW_AVAILABLE = True
except ImportError:
    pass


@dataclass
class DocumentStyle:
    """Document styling configuration"""
    title_font: str = "Arial"
    title_size: int = 24
    heading1_size: int = 16
    heading2_size: int = 14
    heading3_size: int = 12
    body_size: int = 11
    body_font: str = "Times New Roman"
    line_spacing: float = 1.5
    margin_top: float = 2.54  # cm
    margin_bottom: float = 2.54
    margin_left: float = 3.17
    margin_right: float = 2.54
    primary_color: tuple = (0, 51, 102)  # Dark blue
    secondary_color: tuple = (128, 128, 128)  # Gray


class UMLGenerator:
    """Generate UML diagrams using PlantUML or Mermaid"""

    @staticmethod
    def generate_plantuml(uml_code: str, output_path: str, format: str = "png") -> Optional[str]:
        """
        Generate UML diagram using PlantUML

        Args:
            uml_code: PlantUML code
            output_path: Path to save the diagram
            format: Output format (png, svg, eps)

        Returns:
            Path to generated image or None if failed
        """
        try:
            # Create temp file for PlantUML code
            with tempfile.NamedTemporaryFile(mode='w', suffix='.puml', delete=False) as f:
                f.write(uml_code)
                temp_file = f.name

            # Try to run PlantUML
            result = subprocess.run(
                ['plantuml', f'-t{format}', '-o', os.path.dirname(output_path), temp_file],
                capture_output=True,
                timeout=30
            )

            if result.returncode == 0:
                # PlantUML generates file with same name but different extension
                generated_file = temp_file.replace('.puml', f'.{format}')
                if os.path.exists(generated_file):
                    # Move to desired output path
                    os.rename(generated_file, output_path)
                    return output_path

            return None

        except (subprocess.TimeoutExpired, FileNotFoundError):
            return None
        finally:
            # Cleanup temp file
            if os.path.exists(temp_file):
                os.remove(temp_file)

    @staticmethod
    def generate_mermaid(mermaid_code: str, output_path: str) -> Optional[str]:
        """
        Generate diagram using Mermaid CLI

        Args:
            mermaid_code: Mermaid diagram code
            output_path: Path to save the diagram

        Returns:
            Path to generated image or None if failed
        """
        try:
            with tempfile.NamedTemporaryFile(mode='w', suffix='.mmd', delete=False) as f:
                f.write(mermaid_code)
                temp_file = f.name

            result = subprocess.run(
                ['mmdc', '-i', temp_file, '-o', output_path],
                capture_output=True,
                timeout=30
            )

            if result.returncode == 0 and os.path.exists(output_path):
                return output_path

            return None

        except (subprocess.TimeoutExpired, FileNotFoundError):
            return None
        finally:
            if os.path.exists(temp_file):
                os.remove(temp_file)

    @staticmethod
    def get_use_case_diagram(actors: List[str], use_cases: List[str], project_name: str) -> str:
        """Generate PlantUML use case diagram"""
        uml = f"""@startuml
title {project_name} - Use Case Diagram
left to right direction
skinparam packageStyle rectangle

"""
        # Add actors
        for actor in actors:
            uml += f'actor "{actor}" as {actor.replace(" ", "_")}\n'

        uml += "\nrectangle System {\n"

        # Add use cases
        for uc in use_cases:
            uml += f'  usecase "{uc}" as UC_{use_cases.index(uc)}\n'

        uml += "}\n\n"

        # Connect first actor to all use cases (simplified)
        if actors and use_cases:
            actor_id = actors[0].replace(" ", "_")
            for i, _ in enumerate(use_cases):
                uml += f'{actor_id} --> UC_{i}\n'

        uml += "@enduml"
        return uml

    @staticmethod
    def get_class_diagram(classes: List[Dict[str, Any]]) -> str:
        """
        Generate PlantUML class diagram

        Args:
            classes: List of class definitions with name, attributes, methods
        """
        uml = "@startuml\nskinparam classAttributeIconSize 0\n\n"

        for cls in classes:
            uml += f'class {cls["name"]} {{\n'

            # Attributes
            for attr in cls.get("attributes", []):
                uml += f'  {attr}\n'

            uml += "  --\n"

            # Methods
            for method in cls.get("methods", []):
                uml += f'  {method}\n'

            uml += "}\n\n"

        # Add relationships
        for cls in classes:
            for rel in cls.get("relationships", []):
                uml += f'{cls["name"]} {rel["type"]} {rel["target"]}\n'

        uml += "@enduml"
        return uml

    @staticmethod
    def get_sequence_diagram(title: str, participants: List[str], messages: List[Dict]) -> str:
        """Generate PlantUML sequence diagram"""
        uml = f"@startuml\ntitle {title}\n\n"

        for p in participants:
            uml += f'participant "{p}" as {p.replace(" ", "_")}\n'

        uml += "\n"

        for msg in messages:
            from_p = msg["from"].replace(" ", "_")
            to_p = msg["to"].replace(" ", "_")
            uml += f'{from_p} -> {to_p}: {msg["message"]}\n'
            if msg.get("response"):
                uml += f'{to_p} --> {from_p}: {msg["response"]}\n'

        uml += "@enduml"
        return uml

    @staticmethod
    def get_er_diagram(entities: List[Dict[str, Any]]) -> str:
        """Generate PlantUML ER diagram"""
        uml = "@startuml\n!define table(x) class x << (T,#FFAAAA) >>\nhide methods\nhide stereotypes\n\n"

        for entity in entities:
            uml += f'table({entity["name"]}) {{\n'
            for field in entity.get("fields", []):
                pk = " <<PK>>" if field.get("primary_key") else ""
                fk = " <<FK>>" if field.get("foreign_key") else ""
                uml += f'  {field["name"]}: {field["type"]}{pk}{fk}\n'
            uml += "}\n\n"

        # Add relationships
        for entity in entities:
            for rel in entity.get("relationships", []):
                uml += f'{entity["name"]} {rel["type"]} {rel["target"]}: {rel.get("label", "")}\n'

        uml += "@enduml"
        return uml

    @staticmethod
    def get_activity_diagram(title: str, activities: List[str]) -> str:
        """Generate PlantUML activity diagram"""
        uml = f"@startuml\ntitle {title}\nstart\n\n"

        for activity in activities:
            if activity.startswith("if:"):
                condition = activity[3:]
                uml += f'if ({condition}) then (yes)\n'
            elif activity == "else":
                uml += "else (no)\n"
            elif activity == "endif":
                uml += "endif\n"
            else:
                uml += f':{activity};\n'

        uml += "\nstop\n@enduml"
        return uml

    @staticmethod
    def get_deployment_diagram(nodes: List[Dict], project_name: str) -> str:
        """Generate PlantUML deployment diagram"""
        uml = f"@startuml\ntitle {project_name} - Deployment Diagram\n\n"

        for node in nodes:
            uml += f'node "{node["name"]}" {{\n'
            for component in node.get("components", []):
                uml += f'  [{component}]\n'
            uml += "}\n\n"

        uml += "@enduml"
        return uml


class WordDocumentGenerator:
    """Generate Word documents with professional formatting"""

    def __init__(self, style: DocumentStyle = None):
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is required. Install with: pip install python-docx")
        self.style = style or DocumentStyle()
        self.doc = None
        self.images = []

    def create_document(self) -> Document:
        """Create a new Word document with styles"""
        self.doc = Document()
        self._setup_styles()
        self._setup_margins()
        return self.doc

    def _setup_margins(self):
        """Set document margins"""
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Cm(self.style.margin_top)
            section.bottom_margin = Cm(self.style.margin_bottom)
            section.left_margin = Cm(self.style.margin_left)
            section.right_margin = Cm(self.style.margin_right)

    def _setup_styles(self):
        """Setup document styles"""
        styles = self.doc.styles

        # Title style
        if 'CustomTitle' not in [s.name for s in styles]:
            title_style = styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
            title_style.font.name = self.style.title_font
            title_style.font.size = Pt(self.style.title_size)
            title_style.font.bold = True
            title_style.font.color.rgb = RGBColor(*self.style.primary_color)
            title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_style.paragraph_format.space_after = Pt(24)

    def add_cover_page(self, title: str, subtitle: str, team_name: str,
                       college: str, department: str, guide: str,
                       team_members: List[str], academic_year: str,
                       logo_path: Optional[str] = None):
        """Add a professional cover page"""
        # Logo
        if logo_path and os.path.exists(logo_path):
            self.doc.add_picture(logo_path, width=Inches(2))
            self.doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # College name
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(college.upper())
        run.font.size = Pt(16)
        run.font.bold = True

        # Department
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(department)
        run.font.size = Pt(14)

        self.doc.add_paragraph()
        self.doc.add_paragraph()

        # Title
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title.upper())
        run.font.size = Pt(24)
        run.font.bold = True
        run.font.color.rgb = RGBColor(*self.style.primary_color)

        # Subtitle
        if subtitle:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(subtitle)
            run.font.size = Pt(14)
            run.font.italic = True

        self.doc.add_paragraph()
        self.doc.add_paragraph()

        # Team info
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("Submitted by")
        run.font.size = Pt(12)

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(team_name)
        run.font.size = Pt(14)
        run.font.bold = True

        # Team members
        for member in team_members:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(member)
            run.font.size = Pt(11)

        self.doc.add_paragraph()

        # Guide
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("Under the guidance of")
        run.font.size = Pt(12)

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(guide)
        run.font.size = Pt(14)
        run.font.bold = True

        self.doc.add_paragraph()

        # Academic year
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Academic Year: {academic_year}")
        run.font.size = Pt(12)

        # Page break
        self.doc.add_page_break()

    def add_table_of_contents(self):
        """Add table of contents placeholder"""
        p = self.doc.add_paragraph()
        run = p.add_run("TABLE OF CONTENTS")
        run.font.size = Pt(16)
        run.font.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        self.doc.add_paragraph()

        # Add TOC field
        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')

        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'TOC \\o "1-3" \\h \\z \\u'

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')

        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')

        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        run._r.append(fldChar3)

        self.doc.add_paragraph("[Right-click and select 'Update Field' to generate TOC]")
        self.doc.add_page_break()

    def add_heading(self, text: str, level: int = 1):
        """Add a heading"""
        self.doc.add_heading(text, level=level)

    def add_paragraph(self, text: str, bold: bool = False, italic: bool = False):
        """Add a paragraph"""
        p = self.doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = self.style.body_font
        run.font.size = Pt(self.style.body_size)
        run.font.bold = bold
        run.font.italic = italic
        p.paragraph_format.line_spacing = self.style.line_spacing
        return p

    def add_table(self, headers: List[str], rows: List[List[str]],
                  style: str = "Table Grid"):
        """Add a formatted table"""
        table = self.doc.add_table(rows=1, cols=len(headers))
        table.style = style
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Header row
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = header
            # Bold header
            for paragraph in header_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

        # Data rows
        for row_data in rows:
            row_cells = table.add_row().cells
            for i, cell_data in enumerate(row_data):
                row_cells[i].text = str(cell_data)

        self.doc.add_paragraph()  # Space after table
        return table

    def add_image(self, image_path: str, width: float = 5, caption: str = None):
        """Add an image with optional caption"""
        if os.path.exists(image_path):
            self.doc.add_picture(image_path, width=Inches(width))
            self.doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

            if caption:
                p = self.doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(f"Figure: {caption}")
                run.font.size = Pt(10)
                run.font.italic = True

            self.doc.add_paragraph()

    def add_uml_diagram(self, uml_code: str, diagram_type: str, caption: str = None):
        """Add UML diagram generated from code"""
        # Create temp directory for diagrams
        temp_dir = tempfile.mkdtemp()
        output_path = os.path.join(temp_dir, f"{diagram_type}.png")

        # Try to generate diagram
        result = UMLGenerator.generate_plantuml(uml_code, output_path)

        if result and os.path.exists(result):
            self.add_image(result, width=5, caption=caption or f"{diagram_type} Diagram")
            # Cleanup
            os.remove(result)
        else:
            # Fallback: Add code block if diagram generation fails
            self.add_paragraph(f"[{diagram_type} Diagram - PlantUML code below]", italic=True)
            self.add_paragraph(uml_code)

    def add_page_break(self):
        """Add a page break"""
        self.doc.add_page_break()

    def save(self, output_path: str):
        """Save the document"""
        self.doc.save(output_path)
        return output_path


class PDFDocumentGenerator:
    """Generate PDF documents"""

    def __init__(self, style: DocumentStyle = None):
        if not PDF_AVAILABLE:
            raise ImportError("reportlab is required. Install with: pip install reportlab")
        self.style = style or DocumentStyle()
        self.elements = []
        self.styles = getSampleStyleSheet()
        self._setup_styles()

    def _setup_styles(self):
        """Setup custom styles"""
        self.styles.add(ParagraphStyle(
            name='CustomTitle',
            parent=self.styles['Heading1'],
            fontSize=24,
            alignment=1,  # Center
            spaceAfter=30,
            textColor=colors.HexColor('#003366')
        ))

        self.styles.add(ParagraphStyle(
            name='CustomHeading1',
            parent=self.styles['Heading1'],
            fontSize=16,
            spaceBefore=20,
            spaceAfter=10,
            textColor=colors.HexColor('#003366')
        ))

        self.styles.add(ParagraphStyle(
            name='CustomBody',
            parent=self.styles['Normal'],
            fontSize=11,
            leading=16,
            spaceBefore=6,
            spaceAfter=6
        ))

    def add_title(self, text: str):
        """Add title"""
        self.elements.append(Paragraph(text, self.styles['CustomTitle']))
        self.elements.append(Spacer(1, 12))

    def add_heading(self, text: str, level: int = 1):
        """Add heading"""
        style = f'Heading{min(level, 3)}'
        self.elements.append(Paragraph(text, self.styles[style]))
        self.elements.append(Spacer(1, 6))

    def add_paragraph(self, text: str):
        """Add paragraph"""
        self.elements.append(Paragraph(text, self.styles['CustomBody']))

    def add_table(self, data: List[List[str]], col_widths: List[float] = None):
        """Add table"""
        table = Table(data, colWidths=col_widths)
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#003366')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('TEXTCOLOR', (0, 1), (-1, -1), colors.black),
            ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 1), (-1, -1), 10),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        self.elements.append(table)
        self.elements.append(Spacer(1, 12))

    def add_image(self, image_path: str, width: float = 400, caption: str = None):
        """Add image"""
        if os.path.exists(image_path):
            img = Image(image_path, width=width)
            self.elements.append(img)
            if caption:
                self.elements.append(Paragraph(f"<i>Figure: {caption}</i>", self.styles['Normal']))
            self.elements.append(Spacer(1, 12))

    def add_page_break(self):
        """Add page break"""
        self.elements.append(PageBreak())

    def save(self, output_path: str, title: str = "Document"):
        """Save PDF"""
        doc = SimpleDocTemplate(
            output_path,
            pagesize=A4,
            rightMargin=2*cm,
            leftMargin=2.5*cm,
            topMargin=2*cm,
            bottomMargin=2*cm,
            title=title
        )
        doc.build(self.elements)
        return output_path


def check_dependencies() -> Dict[str, bool]:
    """Check which document generation features are available"""
    return {
        "word": DOCX_AVAILABLE,
        "pdf": PDF_AVAILABLE,
        "images": PILLOW_AVAILABLE,
        "plantuml": _check_plantuml(),
        "mermaid": _check_mermaid()
    }


def _check_plantuml() -> bool:
    """Check if PlantUML is available"""
    try:
        result = subprocess.run(['plantuml', '-version'], capture_output=True, timeout=5)
        return result.returncode == 0
    except:
        return False


def _check_mermaid() -> bool:
    """Check if Mermaid CLI is available"""
    try:
        result = subprocess.run(['mmdc', '--version'], capture_output=True, timeout=5)
        return result.returncode == 0
    except:
        return False


def install_dependencies():
    """Return pip install commands for dependencies"""
    return """
# Install document generation dependencies:

# For Word documents:
pip install python-docx

# For PDF documents:
pip install reportlab

# For image processing:
pip install Pillow

# For UML diagrams (PlantUML):
# 1. Install Java (required)
# 2. Download PlantUML: https://plantuml.com/download
# 3. Or use: pip install plantuml

# For Mermaid diagrams:
npm install -g @mermaid-js/mermaid-cli
"""
