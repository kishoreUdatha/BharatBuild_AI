from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# Set margins
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

def add_heading(text, level=1):
    heading = doc.add_heading(text, level=level)
    return heading

def add_paragraph(text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    return p

def add_table(data, header=True):
    table = doc.add_table(rows=len(data), cols=len(data[0]))
    table.style = 'Table Grid'
    for i, row in enumerate(data):
        for j, cell in enumerate(row):
            table.rows[i].cells[j].text = str(cell)
            if header and i == 0:
                table.rows[i].cells[j].paragraphs[0].runs[0].bold = True
    return table

# ===== TITLE =====
title = doc.add_heading("SmartGrow InfoTech - Training Curriculum", 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# Program Overview
add_heading("Program Overview", 1)
add_table([
    ["Parameter", "Details"],
    ["Total Courses", "5"],
    ["Duration per Course", "60-80 Hours"],
    ["Total Program Duration", "300-400 Hours"],
    ["Training Mode", "On-Campus / Hybrid"],
    ["Daily Hours", "4-6 Hours"]
])

doc.add_paragraph()
doc.add_page_break()

# ===== COURSE 1: JAVA =====
add_heading("Course 1: Core Java & Advanced Java", 1)
add_paragraph("Duration: 60-80 Hours", bold=True)

add_heading("Module Breakdown", 2)
add_table([
    ["Module", "Topics", "Hours"],
    ["Module 1: Java Fundamentals", "Introduction, JDK Setup, Variables, Data Types, Operators", "8"],
    ["Module 2: Control Flow & Arrays", "If-else, Switch, Loops, Arrays, Strings", "8"],
    ["Module 3: OOP Concepts", "Classes, Objects, Inheritance, Polymorphism, Abstraction", "12"],
    ["Module 4: Exception Handling", "Try-Catch, Throw, Throws, Custom Exceptions", "6"],
    ["Module 5: Collections Framework", "List, Set, Map, Queue, Iterator, Comparator", "10"],
    ["Module 6: File I/O & Serialization", "File Handling, Streams, Serialization", "6"],
    ["Module 7: JDBC", "Database Connectivity, CRUD Operations, PreparedStatement", "8"],
    ["Module 8: Spring Boot Basics", "Spring Core, REST APIs, Dependency Injection", "10"],
    ["Module 9: Hibernate/JPA", "ORM, Entity Mapping, CRUD with Hibernate", "8"],
    ["Module 10: Project Work", "Mini Project Development", "4"]
])

doc.add_paragraph()
add_heading("Practice Exercises", 2)
add_table([
    ["Week", "Practice Tasks"],
    ["Week 1", "20 Basic Programs (Variables, Loops, Arrays)"],
    ["Week 2", "15 OOP Programs (Classes, Inheritance, Polymorphism)"],
    ["Week 3", "10 Collection Programs + Exception Handling"],
    ["Week 4", "5 JDBC Programs + Spring Boot REST API"],
    ["Week 5", "Mini Project: Student Management System"]
])

doc.add_paragraph()
add_heading("Hands-on Projects", 2)
projects = [
    "Calculator Application - OOP based",
    "Library Management System - Collections + File I/O",
    "Employee Database App - JDBC + MySQL",
    "REST API for E-commerce - Spring Boot + Hibernate"
]
for p in projects:
    doc.add_paragraph(p, style='List Bullet')

doc.add_paragraph()
add_heading("Materials Provided", 2)
materials = [
    "Java Programming eBook (PDF)",
    "Video Tutorials (20+ hours)",
    "Code Examples & Templates",
    "Practice Problem Sets (100+ problems)",
    "Interview Questions Bank (200+ questions)",
    "Project Source Code"
]
for m in materials:
    doc.add_paragraph(m, style='List Bullet')

doc.add_page_break()

# ===== COURSE 2: PYTHON =====
add_heading("Course 2: Python Programming", 1)
add_paragraph("Duration: 60-80 Hours", bold=True)

add_heading("Module Breakdown", 2)
add_table([
    ["Module", "Topics", "Hours"],
    ["Module 1: Python Basics", "Installation, Syntax, Variables, Data Types, Operators", "8"],
    ["Module 2: Control Structures", "If-else, Loops, Break, Continue, Pass", "6"],
    ["Module 3: Data Structures", "Lists, Tuples, Dictionaries, Sets", "10"],
    ["Module 4: Functions & Modules", "Functions, Lambda, Modules, Packages", "8"],
    ["Module 5: File Handling", "Read/Write Files, CSV, JSON, Excel", "6"],
    ["Module 6: OOP in Python", "Classes, Objects, Inheritance, Polymorphism", "8"],
    ["Module 7: Libraries", "NumPy, Pandas, Matplotlib", "10"],
    ["Module 8: Web Framework", "Flask/Django Basics, REST APIs", "10"],
    ["Module 9: Database", "SQLite, MySQL with Python", "6"],
    ["Module 10: Project Work", "Mini Project Development", "8"]
])

doc.add_paragraph()
add_heading("Practice Exercises", 2)
add_table([
    ["Week", "Practice Tasks"],
    ["Week 1", "25 Basic Programs (Variables, Loops, Strings)"],
    ["Week 2", "20 Data Structure Programs (Lists, Dicts)"],
    ["Week 3", "15 Function & OOP Programs"],
    ["Week 4", "10 Library Programs (Pandas, NumPy)"],
    ["Week 5", "Web App Project with Flask/Django"]
])

doc.add_paragraph()
add_heading("Hands-on Projects", 2)
projects = [
    "Number Guessing Game - Basic Python",
    "Contact Book Application - File Handling + OOP",
    "Data Analysis Dashboard - Pandas + Matplotlib",
    "Blog Website - Flask/Django + SQLite",
    "REST API for Todo App - Flask + MySQL"
]
for p in projects:
    doc.add_paragraph(p, style='List Bullet')

doc.add_paragraph()
add_heading("Materials Provided", 2)
materials = [
    "Python Programming eBook (PDF)",
    "Jupyter Notebooks (50+ examples)",
    "Video Tutorials (25+ hours)",
    "Practice Problem Sets (150+ problems)",
    "Cheat Sheets (Pandas, NumPy, Flask)",
    "Interview Questions Bank (150+ questions)",
    "Project Source Code"
]
for m in materials:
    doc.add_paragraph(m, style='List Bullet')

doc.add_page_break()

# ===== COURSE 3: AI/ML =====
add_heading("Course 3: Artificial Intelligence & Machine Learning", 1)
add_paragraph("Duration: 60-80 Hours", bold=True)

add_heading("Module Breakdown", 2)
add_table([
    ["Module", "Topics", "Hours"],
    ["Module 1: AI/ML Introduction", "What is AI/ML, Types, Applications, Tools Setup", "4"],
    ["Module 2: Python for ML", "NumPy, Pandas, Matplotlib, Data Preprocessing", "8"],
    ["Module 3: Statistics & Math", "Probability, Statistics, Linear Algebra Basics", "6"],
    ["Module 4: Supervised Learning", "Linear Regression, Logistic Regression, Decision Trees", "12"],
    ["Module 5: Unsupervised Learning", "K-Means, Hierarchical Clustering, PCA", "8"],
    ["Module 6: Model Evaluation", "Accuracy, Precision, Recall, F1-Score, Cross-Validation", "6"],
    ["Module 7: Deep Learning Basics", "Neural Networks, Activation Functions, Backpropagation", "10"],
    ["Module 8: TensorFlow/Keras", "Building Models, Training, Prediction", "10"],
    ["Module 9: NLP Basics", "Text Processing, Sentiment Analysis", "6"],
    ["Module 10: Projects", "End-to-End ML Projects", "10"]
])

doc.add_paragraph()
add_heading("Practice Exercises", 2)
add_table([
    ["Week", "Practice Tasks"],
    ["Week 1", "Data Preprocessing on 5 Datasets"],
    ["Week 2", "Implement Linear & Logistic Regression from Scratch"],
    ["Week 3", "Classification on 3 Datasets (Decision Tree, Random Forest)"],
    ["Week 4", "Clustering Projects (Customer Segmentation)"],
    ["Week 5", "Neural Network Implementation + Final Project"]
])

doc.add_paragraph()
add_heading("Hands-on Projects", 2)
projects = [
    "House Price Prediction - Linear Regression",
    "Email Spam Classifier - Logistic Regression + NLP",
    "Customer Segmentation - K-Means Clustering",
    "Image Classification - CNN with TensorFlow",
    "Sentiment Analysis - NLP + Deep Learning",
    "Recommendation System - Collaborative Filtering"
]
for p in projects:
    doc.add_paragraph(p, style='List Bullet')

doc.add_paragraph()
add_heading("Materials Provided", 2)
materials = [
    "ML/AI Comprehensive eBook (PDF)",
    "Jupyter Notebooks (30+ notebooks)",
    "Datasets (15+ real-world datasets)",
    "Video Tutorials (30+ hours)",
    "Algorithm Cheat Sheets",
    "TensorFlow/Keras Guide",
    "Interview Questions Bank (100+ questions)",
    "Project Source Code with Documentation"
]
for m in materials:
    doc.add_paragraph(m, style='List Bullet')

doc.add_page_break()

# ===== COURSE 4: REACTJS =====
add_heading("Course 4: ReactJS (Frontend Development)", 1)
add_paragraph("Duration: 60-80 Hours", bold=True)

add_heading("Module Breakdown", 2)
add_table([
    ["Module", "Topics", "Hours"],
    ["Module 1: Web Fundamentals", "HTML5, CSS3, JavaScript ES6+ Refresher", "8"],
    ["Module 2: React Basics", "Create React App, JSX, Components, Props", "10"],
    ["Module 3: State Management", "useState, useEffect, Component Lifecycle", "10"],
    ["Module 4: Advanced Hooks", "useContext, useReducer, useMemo, Custom Hooks", "8"],
    ["Module 5: Routing", "React Router, Navigation, Protected Routes", "6"],
    ["Module 6: Forms & Validation", "Controlled Components, Form Libraries, Validation", "6"],
    ["Module 7: API Integration", "Fetch, Axios, REST API Calls, Error Handling", "8"],
    ["Module 8: State Management", "Redux Toolkit / Context API", "8"],
    ["Module 9: Styling", "CSS Modules, Tailwind CSS, Styled Components", "6"],
    ["Module 10: Project", "Full React Application", "10"]
])

doc.add_paragraph()
add_heading("Practice Exercises", 2)
add_table([
    ["Week", "Practice Tasks"],
    ["Week 1", "10 Component Building Exercises"],
    ["Week 2", "10 State Management Exercises"],
    ["Week 3", "5 Routing + Forms Exercises"],
    ["Week 4", "5 API Integration Projects"],
    ["Week 5", "Full Project Development"]
])

doc.add_paragraph()
add_heading("Hands-on Projects", 2)
projects = [
    "Counter App - useState basics",
    "Todo List - CRUD Operations",
    "Weather App - API Integration",
    "E-commerce Product Page - Routing + State",
    "Blog Platform - Full CRUD + Redux",
    "Dashboard Application - Charts + Tables + Auth"
]
for p in projects:
    doc.add_paragraph(p, style='List Bullet')

doc.add_paragraph()
add_heading("Materials Provided", 2)
materials = [
    "React Complete Guide eBook (PDF)",
    "Code Sandbox Examples (40+)",
    "Video Tutorials (25+ hours)",
    "Component Library Templates",
    "Tailwind CSS Cheat Sheet",
    "Redux Toolkit Guide",
    "Interview Questions Bank (100+ questions)",
    "Project Source Code"
]
for m in materials:
    doc.add_paragraph(m, style='List Bullet')

doc.add_page_break()

# ===== COURSE 5: NODEJS =====
add_heading("Course 5: NodeJS (Backend Development)", 1)
add_paragraph("Duration: 60-80 Hours", bold=True)

add_heading("Module Breakdown", 2)
add_table([
    ["Module", "Topics", "Hours"],
    ["Module 1: Node.js Fundamentals", "Installation, NPM, Modules, Event Loop", "8"],
    ["Module 2: Core Modules", "File System, Path, HTTP, Events", "8"],
    ["Module 3: Express.js", "Routing, Middleware, Error Handling", "10"],
    ["Module 4: REST API Development", "CRUD APIs, Request/Response, Status Codes", "10"],
    ["Module 5: MongoDB", "MongoDB Setup, Mongoose, Schema Design", "10"],
    ["Module 6: Authentication", "JWT, Bcrypt, Passport.js, Sessions", "8"],
    ["Module 7: File Upload & Storage", "Multer, Cloudinary, AWS S3", "4"],
    ["Module 8: Real-time Apps", "Socket.io, WebSockets", "6"],
    ["Module 9: Deployment", "PM2, Docker Basics, Heroku/Railway/Render", "6"],
    ["Module 10: Project", "Full Backend API Project", "10"]
])

doc.add_paragraph()
add_heading("Practice Exercises", 2)
add_table([
    ["Week", "Practice Tasks"],
    ["Week 1", "10 Node.js Core Module Exercises"],
    ["Week 2", "10 Express.js API Exercises"],
    ["Week 3", "5 MongoDB CRUD Projects"],
    ["Week 4", "5 Authentication Implementation Tasks"],
    ["Week 5", "Full Backend Project"]
])

doc.add_paragraph()
add_heading("Hands-on Projects", 2)
projects = [
    "File System CLI Tool - Node.js Core",
    "REST API for Notes - Express + JSON",
    "User Authentication System - JWT + MongoDB",
    "Blog API with Comments - Full CRUD",
    "Real-time Chat Application - Socket.io",
    "E-commerce Backend - Full API + Auth + File Upload"
]
for p in projects:
    doc.add_paragraph(p, style='List Bullet')

doc.add_paragraph()
add_heading("Materials Provided", 2)
materials = [
    "Node.js Complete Guide eBook (PDF)",
    "Code Examples (50+)",
    "Video Tutorials (25+ hours)",
    "API Documentation Templates",
    "Postman Collections",
    "MongoDB Cheat Sheet",
    "Deployment Guides",
    "Interview Questions Bank (100+ questions)",
    "Project Source Code"
]
for m in materials:
    doc.add_paragraph(m, style='List Bullet')

doc.add_page_break()

# ===== ASSESSMENT =====
add_heading("Assessment & Certification", 1)

add_heading("Assessment Pattern", 2)
add_table([
    ["Assessment Type", "Weightage", "Frequency"],
    ["Daily Assignments", "20%", "Daily"],
    ["Weekly Tests", "20%", "Weekly"],
    ["Mini Projects", "30%", "Per Module"],
    ["Final Project", "30%", "End of Course"]
])

doc.add_paragraph()
add_heading("Passing Criteria", 2)
criteria = [
    "Minimum 75% Attendance",
    "Minimum 60% in Assessments",
    "Successful Project Submission"
]
for c in criteria:
    doc.add_paragraph(c, style='List Bullet')

doc.add_paragraph()
add_heading("Certificates Issued", 2)
certs = [
    "Course Completion Certificate - Per Technology",
    "Program Completion Certificate - Overall Program",
    "Project Excellence Certificate - Top Performers"
]
for c in certs:
    doc.add_paragraph(c, style='List Bullet')

doc.add_paragraph()

# ===== MATERIALS SUMMARY =====
add_heading("Materials Summary", 1)

add_heading("Total Program Materials", 2)
add_table([
    ["Material", "Total"],
    ["eBooks", "5"],
    ["Video Hours", "125+ Hours"],
    ["Code Examples", "200+"],
    ["Practice Problems", "600+"],
    ["Interview Questions", "650+"],
    ["Projects", "25+"],
    ["Cheat Sheets", "20+"]
])

doc.add_paragraph()

# ===== DAILY SCHEDULE =====
add_heading("Sample Daily Schedule (6 Hours)", 1)
add_table([
    ["Time", "Activity", "Duration"],
    ["10:00 - 11:30", "Theory Session", "1.5 Hours"],
    ["11:30 - 11:45", "Break", "15 Min"],
    ["11:45 - 13:00", "Hands-on Practice", "1.25 Hours"],
    ["13:00 - 14:00", "Lunch Break", "1 Hour"],
    ["14:00 - 15:30", "Lab/Coding Session", "1.5 Hours"],
    ["15:30 - 15:45", "Break", "15 Min"],
    ["15:45 - 17:00", "Project Work / Q&A", "1.25 Hours"]
])

doc.add_paragraph()
doc.add_paragraph()

# Footer
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("SmartGrow InfoTech Pvt Ltd | January 2026")
run.font.size = Pt(10)
run.italic = True

# Save
output_path = r"C:\Users\Lekhana\Kishore_projects\BharatBuild_AI\docs\SmartGrow_Training_Curriculum.docx"
doc.save(output_path)
print(f"Curriculum saved to: {output_path}")
