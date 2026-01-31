from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Create document
doc = Document()

# Set margins
sections = doc.sections
for section in sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

def add_heading(text, level=1):
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading

def add_centered_heading(text, level=1):
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return heading

def add_paragraph(text, bold=False, italic=False, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    p.alignment = alignment
    return p

def add_table(data, header=True):
    table = doc.add_table(rows=len(data), cols=len(data[0]))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, row in enumerate(data):
        for j, cell in enumerate(row):
            table.rows[i].cells[j].text = str(cell)
            if header and i == 0:
                table.rows[i].cells[j].paragraphs[0].runs[0].bold = True

    return table

def add_signature_line():
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'

    labels = ["Signature:", "Name:", "Designation:", "Date:", "Seal:"]
    for i, label in enumerate(labels):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        table.rows[i].cells[1].text = "_" * 40

    return table

# ===== DOCUMENT CONTENT =====

# Title
title = doc.add_heading("MEMORANDUM OF UNDERSTANDING", 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# Between
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("Between").bold = True

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("SmartGrow InfoTech Pvt Ltd")
run.bold = True
run.font.size = Pt(14)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('(Hereinafter referred to as "SmartGrow" or "First Party")').italic = True

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("And").bold = True

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Scient Institute of Technology")
run.bold = True
run.font.size = Pt(14)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('(Hereinafter referred to as "Scient" or "Second Party")').italic = True

doc.add_paragraph()

# Dates
add_table([
    ["Date of Agreement:", ""],
    ["Effective Date:", ""]
], header=False)

doc.add_paragraph()

# 1. PREAMBLE
add_heading("1. PREAMBLE", level=1)
add_paragraph(
    "This Memorandum of Understanding (MOU) is entered into by and between SmartGrow InfoTech Pvt Ltd "
    "and Scient Institute of Technology for the purpose of establishing a collaborative partnership "
    "for student training and skill development programs."
)
add_paragraph(
    "Both parties recognize the importance of bridging the gap between academic learning and industry "
    "requirements, and hereby agree to work together to enhance the employability of students through "
    "practical, industry-oriented training programs."
)

# 2. PARTIES
add_heading("2. PARTIES TO THE AGREEMENT", level=1)

add_heading("2.1 First Party", level=2)
add_paragraph("SmartGrow InfoTech Pvt Ltd", bold=True)
add_paragraph("Address: _____________________________________________")
add_paragraph("Represented by: _________________ (Designation: _____________)")
add_paragraph("Contact: _________________")
add_paragraph("Email: _________________")

add_heading("2.2 Second Party", level=2)
add_paragraph("Scient Institute of Technology", bold=True)
add_paragraph("Address: _____________________________________________")
add_paragraph("Represented by: _________________ (Designation: _____________)")
add_paragraph("Contact: _________________")
add_paragraph("Email: _________________")

# 3. OBJECTIVES
add_heading("3. OBJECTIVES", level=1)
add_paragraph("The objectives of this MOU are:")
objectives = [
    "To provide industry-relevant technical training to students of Scient Institute of Technology.",
    "To enhance students' practical skills and employability through hands-on project-based learning.",
    "To bridge the gap between academic curriculum and industry requirements.",
    "To create a sustainable partnership for continuous skill development initiatives.",
    "To provide exposure to emerging technologies and industry best practices."
]
for i, obj in enumerate(objectives, 1):
    add_paragraph(f"3.{i} {obj}")

# 4. SCOPE OF TRAINING
add_heading("4. SCOPE OF TRAINING PROGRAM", level=1)

add_heading("4.1 Training Technologies", level=2)
add_paragraph("SmartGrow shall provide training in the following technologies:")
doc.add_paragraph()
add_table([
    ["S.No", "Technology", "Duration"],
    ["1", "Core Java & Advanced Java", "60-80 Hours"],
    ["2", "Python Programming", "60-80 Hours"],
    ["3", "Artificial Intelligence & Machine Learning", "60-80 Hours"],
    ["4", "ReactJS (Frontend Development)", "60-80 Hours"],
    ["5", "NodeJS (Backend Development)", "60-80 Hours"]
])
doc.add_paragraph()

add_heading("4.2 Training Details", level=2)
add_table([
    ["Parameter", "Details"],
    ["Duration per Course", "60-80 Hours"],
    ["Total Program Duration", "300-400 Hours (5 Courses)"],
    ["Mode", "On-Campus / Hybrid"],
    ["Batch Size", "As mutually agreed"],
    ["Number of Students", "400 (Approximate)"],
    ["Training Hours", "4-6 Hours per day"]
], header=True)
doc.add_paragraph()

add_heading("4.3 Training Deliverables", level=2)
deliverables = [
    "Comprehensive course material and resources",
    "Hands-on practical sessions with real-world projects",
    "Assessment tests and evaluations",
    "Certificates of completion",
    "Placement assistance and guidance"
]
for d in deliverables:
    doc.add_paragraph(d, style='List Bullet')

# 5. RESPONSIBILITIES OF SMARTGROW
add_heading("5. RESPONSIBILITIES OF SMARTGROW (First Party)", level=1)
add_paragraph("SmartGrow shall be responsible for:")
smartgrow_resp = [
    "Providing qualified and experienced trainers/faculty for the training program.",
    "Developing and delivering comprehensive training curriculum and materials.",
    "Conducting practical sessions, assignments, and assessments.",
    "Providing necessary software, tools, and cloud resources for training.",
    "Issuing certificates of completion to students who successfully complete the program.",
    "Providing placement assistance and industry connect opportunities.",
    "Maintaining training quality standards and providing progress reports.",
    "Addressing queries and providing support during the training period."
]
for i, resp in enumerate(smartgrow_resp, 1):
    add_paragraph(f"5.{i} {resp}")

# 6. RESPONSIBILITIES OF SCIENT
add_heading("6. RESPONSIBILITIES OF SCIENT (Second Party)", level=1)
add_paragraph("Scient Institute shall be responsible for:")
scient_resp = [
    "Providing suitable infrastructure including classrooms, computer labs, projectors, and internet connectivity.",
    "Ensuring student attendance and discipline during training sessions.",
    "Appointing a coordinator to liaise with SmartGrow for smooth execution.",
    "Providing student data and batch schedules in advance.",
    "Timely payment of training fees as per agreed terms.",
    "Promoting the training program among students and encouraging participation.",
    "Providing necessary administrative support for training execution.",
    "Ensuring availability of power backup and technical infrastructure."
]
for i, resp in enumerate(scient_resp, 1):
    add_paragraph(f"6.{i} {resp}")

# 7. FINANCIAL TERMS
add_heading("7. FINANCIAL TERMS", level=1)

add_heading("7.1 Training Fee Structure", level=2)
add_table([
    ["Particulars", "Amount"],
    ["Training Fee per Student", "Rs. 14,000/- (Rupees Fourteen Thousand Only)"],
    ["Total Students (Estimated)", "400"],
    ["Total Training Value", "Rs. 56,00,000/- (Rupees Fifty-Six Lakhs Only)"]
])
doc.add_paragraph()

add_heading("7.2 Payment Terms", level=2)
add_table([
    ["Milestone", "Percentage", "Amount", "Due Date"],
    ["Upon Signing MOU", "30%", "Rs. 16,80,000/-", "Within 7 days of signing"],
    ["Mid-Training", "40%", "Rs. 22,40,000/-", "After 50% training completion"],
    ["Upon Completion", "30%", "Rs. 16,80,000/-", "Within 7 days of completion"]
])
doc.add_paragraph()

add_heading("7.3 Payment Mode", level=2)
add_paragraph("All payments shall be made via Bank Transfer / Cheque / DD")
add_paragraph("Payments to be made in favor of: SmartGrow InfoTech Pvt Ltd", bold=True)
add_paragraph("GST as applicable shall be charged additionally")

add_heading("7.4 Bank Details", level=2)
add_table([
    ["Account Name", "SmartGrow InfoTech Pvt Ltd"],
    ["Bank Name", ""],
    ["Account Number", ""],
    ["IFSC Code", ""],
    ["Branch", ""]
], header=False)
doc.add_paragraph()

# 8. DURATION
add_heading("8. DURATION OF MOU", level=1)
add_paragraph("8.1 This MOU shall be effective from the date of signing and shall remain valid for a period of One (1) Year.")
add_paragraph("8.2 The MOU may be renewed for subsequent years by mutual written consent of both parties.")
add_paragraph("8.3 Either party may terminate this MOU by providing 30 days written notice to the other party.")

# 9. CONFIDENTIALITY
add_heading("9. CONFIDENTIALITY", level=1)
add_paragraph("9.1 Both parties agree to maintain confidentiality of all proprietary information, training materials, student data, and business information shared during the course of this partnership.")
add_paragraph("9.2 Neither party shall disclose confidential information to third parties without prior written consent.")
add_paragraph("9.3 This confidentiality obligation shall survive the termination of this MOU.")

# 10. INTELLECTUAL PROPERTY
add_heading("10. INTELLECTUAL PROPERTY", level=1)
add_paragraph("10.1 All training materials, course content, software, and tools developed or provided by SmartGrow shall remain the intellectual property of SmartGrow.")
add_paragraph("10.2 Scient Institute shall not reproduce, distribute, or share SmartGrow's proprietary content without prior written approval.")
add_paragraph("10.3 Projects developed by students during training shall be jointly owned for academic and portfolio purposes.")

# 11. CERTIFICATION
add_heading("11. CERTIFICATION", level=1)
add_paragraph("11.1 Students who successfully complete the training program with minimum 75% attendance and passing assessment scores shall receive:")
certs = [
    "Certificate of Completion from SmartGrow InfoTech",
    "Technology-specific skill certificates",
    "Project completion certificates"
]
for c in certs:
    doc.add_paragraph(c, style='List Bullet')
add_paragraph("11.2 Certificates shall be issued within 15 working days of program completion.")

# 12. PLACEMENT ASSISTANCE
add_heading("12. PLACEMENT ASSISTANCE", level=1)
add_paragraph("12.1 SmartGrow shall provide placement assistance to eligible students including:")
placement = [
    "Resume building and interview preparation",
    "Mock interviews and soft skills training",
    "Industry connect and job referrals",
    "Campus placement coordination support"
]
for p in placement:
    doc.add_paragraph(p, style='List Bullet')
add_paragraph("12.2 Placement assistance is provided on best-effort basis and does not guarantee employment.")

# 13. TERMINATION
add_heading("13. TERMINATION", level=1)
add_paragraph("13.1 Either party may terminate this MOU by providing 30 days written notice.")
add_paragraph("13.2 In case of termination:")
termination = [
    "Payments for completed training shall be settled",
    "Pending batches shall be completed or mutually adjusted",
    "All confidential materials shall be returned"
]
for t in termination:
    doc.add_paragraph(t, style='List Bullet')
add_paragraph("13.3 Termination due to breach of terms shall be immediate upon written notice.")

# 14. DISPUTE RESOLUTION
add_heading("14. DISPUTE RESOLUTION", level=1)
add_paragraph("14.1 Any disputes arising from this MOU shall first be resolved through mutual discussion and negotiation.")
add_paragraph("14.2 If disputes cannot be resolved amicably, they shall be referred to arbitration as per the Arbitration and Conciliation Act, 1996.")
add_paragraph("14.3 The jurisdiction for any legal proceedings shall be _________________ (City).")

# 15. FORCE MAJEURE
add_heading("15. FORCE MAJEURE", level=1)
add_paragraph(
    "Neither party shall be liable for failure to perform obligations due to circumstances beyond their control, "
    "including but not limited to natural disasters, pandemics, government actions, or civil unrest. In such cases, "
    "the affected party shall notify the other party promptly, and both parties shall work together to resume "
    "obligations as soon as practicable."
)

# 16. AMENDMENTS
add_heading("16. AMENDMENTS", level=1)
add_paragraph("Any amendments or modifications to this MOU shall be made in writing and signed by authorized representatives of both parties.")

# 17. GENERAL PROVISIONS
add_heading("17. GENERAL PROVISIONS", level=1)
add_paragraph("17.1 This MOU represents the entire agreement between the parties and supersedes all prior negotiations and agreements.")
add_paragraph("17.2 Neither party shall assign or transfer this MOU without prior written consent of the other party.")
add_paragraph("17.3 The failure of either party to enforce any provision shall not constitute a waiver of that provision.")
add_paragraph("17.4 If any provision is found invalid, the remaining provisions shall continue in full force.")

# 18. SIGNATURES
doc.add_page_break()
add_heading("18. SIGNATURES", level=1)
add_paragraph(
    "IN WITNESS WHEREOF, the parties have executed this Memorandum of Understanding as of the date first written above.",
    bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER
)

doc.add_paragraph()
doc.add_paragraph()

# First Party Signature
add_heading("FOR SMARTGROW INFOTECH PVT LTD (First Party)", level=2)
add_signature_line()

doc.add_paragraph()
doc.add_paragraph()

# Second Party Signature
add_heading("FOR SCIENT INSTITUTE OF TECHNOLOGY (Second Party)", level=2)
add_signature_line()

doc.add_paragraph()
doc.add_paragraph()

# Witnesses
add_heading("WITNESSES", level=2)

add_paragraph("Witness 1:", bold=True)
add_paragraph("Signature: _________________________________")
add_paragraph("Name: _________________________________")
add_paragraph("Address: _________________________________")

doc.add_paragraph()

add_paragraph("Witness 2:", bold=True)
add_paragraph("Signature: _________________________________")
add_paragraph("Name: _________________________________")
add_paragraph("Address: _________________________________")

# ANNEXURES
doc.add_page_break()
add_heading("ANNEXURE A: TRAINING SCHEDULE", level=1)
add_table([
    ["S.No", "Technology", "Topics Covered", "Hours"],
    ["1", "Core Java & Advanced Java", "OOPs, Collections, JDBC, Spring Boot, Hibernate", "60-80"],
    ["2", "Python Programming", "Basics, Data Structures, Libraries, Django/Flask", "60-80"],
    ["3", "AI & Machine Learning", "ML Algorithms, Neural Networks, TensorFlow, Projects", "60-80"],
    ["4", "ReactJS", "Components, Hooks, State Management, REST APIs", "60-80"],
    ["5", "NodeJS", "Express, MongoDB, Authentication, Deployment", "60-80"]
])
doc.add_paragraph()
add_paragraph("Total Program Duration: 300-400 Hours", bold=True)

doc.add_paragraph()
doc.add_paragraph()

add_heading("ANNEXURE B: INFRASTRUCTURE REQUIREMENTS", level=1)
add_table([
    ["Requirement", "Specification", "Responsibility"],
    ["Computer Lab", "50+ systems with internet", "Scient"],
    ["Projector & Screen", "HD Display", "Scient"],
    ["Internet", "Minimum 50 Mbps", "Scient"],
    ["Power Backup", "UPS / Generator", "Scient"],
    ["Software & Tools", "IDEs, Cloud Access", "SmartGrow"],
    ["Training Materials", "Digital & Print", "SmartGrow"]
])

doc.add_paragraph()
doc.add_paragraph()

add_heading("ANNEXURE C: CONTACT PERSONS", level=1)

add_paragraph("SmartGrow InfoTech:", bold=True)
add_paragraph("Program Manager: _____________________")
add_paragraph("Phone: _____________________")
add_paragraph("Email: _____________________")

doc.add_paragraph()

add_paragraph("Scient Institute:", bold=True)
add_paragraph("Coordinator: _____________________")
add_paragraph("Phone: _____________________")
add_paragraph("Email: _____________________")

doc.add_paragraph()
doc.add_paragraph()

# Footer
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("This MOU is executed in duplicate, with each party retaining one original copy.").italic = True

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p.add_run("Document Version: 1.0 | January 2026").font.size = Pt(9)

# Save
output_path = r"C:\Users\Lekhana\Kishore_projects\BharatBuild_AI\docs\MOU_SmartGrow_Scient_Updated.docx"
doc.save(output_path)
print(f"MOU Word document saved to: {output_path}")
