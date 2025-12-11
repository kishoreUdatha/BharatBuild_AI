"""
WORD DOCUMENT GENERATOR
=======================
Creates beautiful 60-80 page Word documents with professional formatting.

Features:
- Professional academic styling
- Auto-generated Table of Contents
- Proper chapter/section numbering
- Figures and tables with captions
- Code blocks with syntax highlighting
- Headers and footers
- Page numbers
"""

from typing import Dict, List, Optional, Any
from datetime import datetime
import os
import tempfile
import re

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from app.core.logging_config import logger
from app.core.config import settings
from app.modules.automation.uml_generator import uml_generator


class WordDocumentGenerator:
    """
    Professional Word Document Generator for academic projects

    Creates beautiful documents with:
    - Custom styles
    - Proper formatting
    - Auto TOC
    - Figures/Tables
    """

    # Color scheme
    PRIMARY_COLOR = RGBColor(0, 51, 102)  # Dark blue
    SECONDARY_COLOR = RGBColor(51, 51, 51)  # Dark gray
    ACCENT_COLOR = RGBColor(0, 102, 153)  # Teal

    def __init__(self):
        self.document = None
        self.styles_created = False

    async def create_document(
        self,
        sections: List[Dict],
        project_data: Dict,
        document_type: str,
        project_id: str = None,
        user_id: str = None
    ) -> str:
        """
        Create complete Word document from sections.

        Args:
            sections: List of generated section content
            project_data: Project metadata
            document_type: Type of document (project_report, srs, etc.)
            project_id: Project ID for saving to project's docs folder
            user_id: User ID for isolation

        Returns:
            Path to generated document
        """
        # Store user_id for diagram generation
        self.user_id = user_id
        self.project_id = project_id

        try:
            # Create document
            self.document = Document()

            # Setup document properties
            self._setup_document_properties(project_data)

            # Create custom styles
            self._create_custom_styles()

            # Setup page layout
            self._setup_page_layout()

            # Add sections
            for section in sections:
                await self._add_section(section, project_data)

            # Add table of contents (placeholder - Word will update)
            self._update_toc_fields()

            # Determine output directory - prefer project docs folder with user isolation
            if project_id:
                output_dir = settings.get_project_docs_dir(project_id, user_id)
            else:
                output_dir = settings.GENERATED_DIR / "documents"
            output_dir.mkdir(parents=True, exist_ok=True)

            filename = f"{project_data.get('project_name', 'Document')}_{document_type}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            file_path = output_dir / filename

            self.document.save(str(file_path))

            logger.info(f"[WordGenerator] Created document: {file_path}")

            return str(file_path)

        except Exception as e:
            logger.error(f"[WordGenerator] Error: {e}", exc_info=True)
            raise

    def _setup_document_properties(self, project_data: Dict):
        """Set document properties"""
        core_props = self.document.core_properties
        core_props.title = project_data.get("project_name", "Project Document")
        core_props.author = project_data.get("author", "BharatBuild AI")
        core_props.subject = project_data.get("project_type", "Software Project")
        core_props.keywords = "project, documentation, academic"

    def _setup_page_layout(self):
        """Setup page layout - margins, headers, footers"""
        section = self.document.sections[0]

        # A4 page size
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)

        # Margins
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)  # Binding margin
        section.right_margin = Cm(2.54)

        # Header and footer
        section.header_distance = Cm(1.27)
        section.footer_distance = Cm(1.27)

        # Add page numbers to footer
        self._add_page_numbers(section)

    def _add_page_numbers(self, section):
        """Add page numbers to footer with proper formatting"""
        footer = section.footer
        footer.is_linked_to_previous = False

        # Clear existing footer content
        for para in footer.paragraphs:
            para.clear()

        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add "Page " text
        page_text_run = paragraph.add_run("Page ")
        page_text_run.font.size = Pt(10)
        page_text_run.font.name = 'Times New Roman'

        # Add page number field
        run = paragraph.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        run._r.append(fldChar1)

        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = ' PAGE '
        run._r.append(instrText)

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        run._r.append(fldChar2)

        # Add placeholder number (will be replaced by Word)
        num_run = paragraph.add_run("1")
        num_run.font.size = Pt(10)
        num_run.font.name = 'Times New Roman'

        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar3)

        # Add " of " text
        of_text_run = paragraph.add_run(" of ")
        of_text_run.font.size = Pt(10)
        of_text_run.font.name = 'Times New Roman'

        # Add total pages field
        run2 = paragraph.add_run()
        fldChar4 = OxmlElement('w:fldChar')
        fldChar4.set(qn('w:fldCharType'), 'begin')
        run2._r.append(fldChar4)

        instrText2 = OxmlElement('w:instrText')
        instrText2.set(qn('xml:space'), 'preserve')
        instrText2.text = ' NUMPAGES '
        run2._r.append(instrText2)

        fldChar5 = OxmlElement('w:fldChar')
        fldChar5.set(qn('w:fldCharType'), 'separate')
        run2._r.append(fldChar5)

        # Add placeholder total (will be replaced by Word)
        total_run = paragraph.add_run("1")
        total_run.font.size = Pt(10)
        total_run.font.name = 'Times New Roman'

        fldChar6 = OxmlElement('w:fldChar')
        fldChar6.set(qn('w:fldCharType'), 'end')
        run2._r.append(fldChar6)

    def _create_custom_styles(self):
        """Create custom styles for professional look"""
        if self.styles_created:
            return

        styles = self.document.styles

        # Title Style
        try:
            title_style = styles.add_style('ProjectTitle', WD_STYLE_TYPE.PARAGRAPH)
            title_style.font.name = 'Arial'
            title_style.font.size = Pt(28)
            title_style.font.bold = True
            title_style.font.color.rgb = self.PRIMARY_COLOR
            title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_style.paragraph_format.space_after = Pt(24)
        except:
            pass

        # Chapter Heading Style
        try:
            chapter_style = styles.add_style('ChapterHeading', WD_STYLE_TYPE.PARAGRAPH)
            chapter_style.font.name = 'Arial'
            chapter_style.font.size = Pt(18)
            chapter_style.font.bold = True
            chapter_style.font.color.rgb = self.PRIMARY_COLOR
            chapter_style.paragraph_format.space_before = Pt(24)
            chapter_style.paragraph_format.space_after = Pt(12)
            chapter_style.paragraph_format.page_break_before = True
        except:
            pass

        # Section Heading Style
        try:
            section_style = styles.add_style('SectionHeading', WD_STYLE_TYPE.PARAGRAPH)
            section_style.font.name = 'Arial'
            section_style.font.size = Pt(14)
            section_style.font.bold = True
            section_style.font.color.rgb = self.SECONDARY_COLOR
            section_style.paragraph_format.space_before = Pt(18)
            section_style.paragraph_format.space_after = Pt(8)
        except:
            pass

        # Subsection Heading Style
        try:
            subsection_style = styles.add_style('SubsectionHeading', WD_STYLE_TYPE.PARAGRAPH)
            subsection_style.font.name = 'Arial'
            subsection_style.font.size = Pt(12)
            subsection_style.font.bold = True
            subsection_style.font.color.rgb = self.SECONDARY_COLOR
            subsection_style.paragraph_format.space_before = Pt(12)
            subsection_style.paragraph_format.space_after = Pt(6)
        except:
            pass

        # Body Text Style - MUST be 12pt with 1.5 line spacing
        try:
            body_style = styles.add_style('BodyText', WD_STYLE_TYPE.PARAGRAPH)
            body_style.font.name = 'Times New Roman'
            body_style.font.size = Pt(12)
            body_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            body_style.paragraph_format.space_after = Pt(8)
            body_style.paragraph_format.first_line_indent = Cm(1.27)
        except:
            pass

        # Also modify the default Normal style to ensure consistency
        try:
            normal_style = styles['Normal']
            normal_style.font.name = 'Times New Roman'
            normal_style.font.size = Pt(12)
            normal_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        except:
            pass

        # Modify List Bullet style
        try:
            if 'List Bullet' in styles:
                list_style = styles['List Bullet']
                list_style.font.name = 'Times New Roman'
                list_style.font.size = Pt(12)
                list_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        except:
            pass

        # Modify List Number style
        try:
            if 'List Number' in styles:
                list_num_style = styles['List Number']
                list_num_style.font.name = 'Times New Roman'
                list_num_style.font.size = Pt(12)
                list_num_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        except:
            pass

        # Code Style
        try:
            code_style = styles.add_style('CodeBlock', WD_STYLE_TYPE.PARAGRAPH)
            code_style.font.name = 'Consolas'
            code_style.font.size = Pt(10)
            code_style.paragraph_format.space_before = Pt(6)
            code_style.paragraph_format.space_after = Pt(6)
        except:
            pass

        # Caption Style
        try:
            caption_style = styles.add_style('FigureCaption', WD_STYLE_TYPE.PARAGRAPH)
            caption_style.font.name = 'Arial'
            caption_style.font.size = Pt(10)
            caption_style.font.italic = True
            caption_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption_style.paragraph_format.space_before = Pt(6)
            caption_style.paragraph_format.space_after = Pt(12)
        except:
            pass

        self.styles_created = True

    async def _add_section(self, section: Dict, project_data: Dict):
        """Add a section to the document"""
        section_type = section.get("type", "generate")
        section_id = section.get("section_id", "")
        title = section.get("title", "")
        content = section.get("content", {})

        if section_type == "template":
            self._add_template_section(section_id, content, project_data)
        elif section_type == "auto":
            self._add_auto_section(section_id)
        elif section_type == "code":
            self._add_code_section(content)
        else:
            self._add_generated_section(section_id, title, content, project_data)

    def _add_template_section(self, section_id: str, content: Dict, project_data: Dict):
        """Add template-based section (cover, certificate, etc.)"""

        if section_id == "cover":
            self._add_cover_page(content, project_data)
        elif section_id == "certificate":
            self._add_certificate_page(content, project_data)
        elif section_id == "declaration":
            self._add_declaration_page(content, project_data)
        elif section_id == "acknowledgement":
            self._add_acknowledgement_page(content, project_data)

    def _add_cover_page(self, content: Dict, project_data: Dict):
        """Add cover page with full college information"""
        # College name
        p = self.document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(content.get("college_name", project_data.get("institution", "University Name")))
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = self.PRIMARY_COLOR

        # Affiliated to
        if content.get("affiliated_to"):
            p = self.document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f"({content.get('affiliated_to')})")
            run.font.size = Pt(12)
            run.font.italic = True

        # College address
        if content.get("college_address"):
            p = self.document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(content.get("college_address"))
            run.font.size = Pt(11)

        self.document.add_paragraph()

        # Department
        p = self.document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(content.get("department", project_data.get("department", "Department of Computer Science")))
        run.font.size = Pt(14)
        run.font.bold = True

        # Add spacing
        self.document.add_paragraph()
        self.document.add_paragraph()

        # Project Report header
        p = self.document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("PROJECT REPORT")
        run.font.size = Pt(14)
        run.font.bold = True

        p = self.document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("On")
        run.font.size = Pt(12)

        # Project Title
        self.document.add_paragraph()
        p = self.document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f'"{content.get("project_name", project_data.get("project_name", "Project Title"))}"')
        run.font.size = Pt(20)
        run.font.bold = True
        run.font.color.rgb = self.PRIMARY_COLOR

        # Subtitle
        p = self.document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("Submitted in partial fulfillment of the requirements for the award of")
        run.font.size = Pt(11)

        p = self.document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("BACHELOR OF TECHNOLOGY")
        run.font.size = Pt(14)
        run.font.bold = True

        p = self.document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("In")
        run.font.size = Pt(11)

        p = self.document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("COMPUTER SCIENCE AND ENGINEERING")
        run.font.size = Pt(12)
        run.font.bold = True

        # Add spacing
        self.document.add_paragraph()

        # Submitted by - Student table
        p = self.document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("Submitted by:")
        run.font.size = Pt(12)
        run.font.bold = True

        students = content.get("students", [])
        if students:
            # Create student table
            table = self.document.add_table(rows=len(students) + 1, cols=3)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            # Header row
            header_cells = table.rows[0].cells
            header_cells[0].text = "S.No"
            header_cells[1].text = "Name"
            header_cells[2].text = "Roll Number"
            for cell in header_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

            # Student rows
            for idx, student in enumerate(students):
                row = table.rows[idx + 1]
                row.cells[0].text = str(idx + 1)
                row.cells[1].text = student.get("name", "Student")
                row.cells[2].text = student.get("roll_number", "")

        self.document.add_paragraph()

        # Guide
        p = self.document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("Under the guidance of")
        run.font.size = Pt(12)
        run.font.bold = True

        p = self.document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(content.get("guide_name", project_data.get("guide", "Guide Name")))
        run.font.size = Pt(12)

        # Academic year
        self.document.add_paragraph()
        p = self.document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Academic Year: {content.get('academic_year', content.get('date', datetime.now().strftime('%Y-%Y')))}")
        run.font.size = Pt(12)
        run.font.bold = True

        # Page break
        self.document.add_page_break()

    def _add_certificate_page(self, content: Dict, project_data: Dict):
        """Add certificate page with full college information"""
        # College header
        college_name = content.get("college_name", project_data.get("institution", "University Name"))
        p = self.document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(college_name)
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = self.PRIMARY_COLOR

        # Affiliated to
        if content.get("affiliated_to"):
            p = self.document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f"({content.get('affiliated_to')})")
            run.font.size = Pt(11)
            run.font.italic = True

        # Department
        p = self.document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(content.get("department", "Department of Computer Science and Engineering"))
        run.font.size = Pt(12)
        run.font.bold = True

        self.document.add_paragraph()

        # Certificate title
        p = self.document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("CERTIFICATE")
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.underline = True

        self.document.add_paragraph()

        project_title = content.get("project_title", project_data.get("project_name", "Project"))
        academic_year = content.get("academic_year", "2024-2025")

        cert_text = f"""This is to certify that the project entitled "{project_title}" is a bonafide work carried out by the following students in partial fulfillment of the requirements for the award of Bachelor of Technology in Computer Science and Engineering from {college_name} during the academic year {academic_year}.

This project work has been approved as it satisfies the academic requirements prescribed for the said degree."""

        p = self.document.add_paragraph(cert_text)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        p.paragraph_format.space_after = Pt(12)

        # Student table
        students = content.get("students", [])
        if students:
            table = self.document.add_table(rows=len(students) + 1, cols=3)
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            # Header
            header_cells = table.rows[0].cells
            header_cells[0].text = "S.No"
            header_cells[1].text = "Name"
            header_cells[2].text = "Roll Number"
            for cell in header_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

            # Students
            for idx, student in enumerate(students):
                row = table.rows[idx + 1]
                row.cells[0].text = str(idx + 1)
                row.cells[1].text = student.get("name", "Student")
                row.cells[2].text = student.get("roll_number", "")

        self.document.add_paragraph()
        self.document.add_paragraph()

        # Signature table - 3 columns
        sig_table = self.document.add_table(rows=4, cols=3)
        sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Row 1: Titles
        sig_table.cell(0, 0).text = "Project Guide"
        sig_table.cell(0, 1).text = "Head of Department"
        sig_table.cell(0, 2).text = "Principal"

        # Row 2: Names
        sig_table.cell(1, 0).text = content.get("guide_name", project_data.get("guide", "Guide Name"))
        sig_table.cell(1, 1).text = content.get("hod_name", "HOD Name")
        sig_table.cell(1, 2).text = content.get("principal_name", "Principal Name")

        # Row 3: Signature line
        sig_table.cell(2, 0).text = "Signature: ____________"
        sig_table.cell(2, 1).text = "Signature: ____________"
        sig_table.cell(2, 2).text = "Signature: ____________"

        # Row 4: Date
        sig_table.cell(3, 0).text = "Date: ____________"
        sig_table.cell(3, 1).text = "Date: ____________"
        sig_table.cell(3, 2).text = "Date: ____________"

        # Bold titles
        for cell in sig_table.rows[0].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        self.document.add_paragraph()

        # External Examiner
        p = self.document.add_paragraph()
        run = p.add_run("External Examiner")
        run.font.bold = True

        self.document.add_paragraph("Name: ________________________")
        self.document.add_paragraph("Signature: ____________________")
        self.document.add_paragraph("Date: ________________________")

        self.document.add_page_break()

    def _add_declaration_page(self, content: Dict, project_data: Dict):
        """Add declaration page with student signatures"""
        p = self.document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("DECLARATION")
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.underline = True

        self.document.add_paragraph()

        project_title = content.get("project_title", project_data.get("project_name", "Project"))
        college_name = content.get("college_name", project_data.get("institution", "University"))
        department = content.get("department", "Department of Computer Science and Engineering")
        guide_name = content.get("guide_name", project_data.get("guide", "Guide Name"))

        decl_text = f"""We, the undersigned, hereby declare that the project entitled "{project_title}" submitted to {college_name}, {department}, is a record of an original work done by us under the guidance of {guide_name}.

This project work is submitted in partial fulfillment of the requirements for the award of the degree of Bachelor of Technology in Computer Science and Engineering.

We further declare that:"""

        p = self.document.add_paragraph(decl_text)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        p.paragraph_format.space_after = Pt(12)

        # Declaration points
        points = [
            "This project is based on our original work.",
            "This project has not been submitted previously for any degree or examination in any other university.",
            "All sources of information have been duly acknowledged.",
            "We have followed the guidelines provided by the institute for preparing this report."
        ]

        for idx, point in enumerate(points, 1):
            p = self.document.add_paragraph(f"{idx}. {point}")
            p.paragraph_format.left_indent = Cm(1)
            p.paragraph_format.space_after = Pt(6)

        self.document.add_paragraph()

        # Student signatures
        p = self.document.add_paragraph()
        run = p.add_run("Student Signatures:")
        run.font.bold = True

        self.document.add_paragraph()

        students = content.get("students", [])
        if students:
            # Create signature table
            sig_table = self.document.add_table(rows=len(students) + 1, cols=3)
            sig_table.style = 'Table Grid'
            sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER

            # Header
            header_cells = sig_table.rows[0].cells
            header_cells[0].text = "Name"
            header_cells[1].text = "Roll Number"
            header_cells[2].text = "Signature"
            for cell in header_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

            # Student rows
            for idx, student in enumerate(students):
                row = sig_table.rows[idx + 1]
                row.cells[0].text = student.get("name", "Student")
                row.cells[1].text = student.get("roll_number", "")
                row.cells[2].text = "________________"

        self.document.add_paragraph()
        self.document.add_paragraph()

        # Date and Place
        date_str = content.get("date", datetime.now().strftime("%B %Y"))
        p = self.document.add_paragraph()
        p.add_run(f"Date: {date_str}")

        p = self.document.add_paragraph()
        p.add_run(f"Place: {college_name}")

        self.document.add_page_break()

    def _add_acknowledgement_page(self, content: Dict, project_data: Dict):
        """Add acknowledgement page"""
        p = self.document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("ACKNOWLEDGEMENT")
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.underline = True

        self.document.add_paragraph()

        guide_name = content.get("guide_name", project_data.get("guide", "our project guide"))
        hod_name = content.get("hod_name", "the Head of Department")
        principal_name = content.get("principal_name", "the Principal")
        college_name = content.get("college_name", project_data.get("institution", "the college"))
        department = content.get("department", "the department")

        paragraphs = [
            f"We take this opportunity to express our profound gratitude and deep regards to our project guide {guide_name} for the exemplary guidance, monitoring, and constant encouragement throughout the course of this project.",
            f"We would like to express our sincere thanks to {hod_name}, Head of Department, {department}, for providing us with the opportunity to work on this project.",
            f"We also express our sincere gratitude to {principal_name}, Principal, {college_name}, for providing us with the necessary facilities and support.",
            f"We extend our heartfelt thanks to all the faculty members of the {department} for their valuable suggestions and support during the development of this project.",
            "We would also like to thank our family and friends for their constant support and encouragement.",
            "Finally, we thank all those who directly or indirectly helped us in the successful completion of this project."
        ]

        for para_text in paragraphs:
            p = self.document.add_paragraph(para_text)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            p.paragraph_format.space_after = Pt(12)
            p.paragraph_format.first_line_indent = Cm(1.27)

        self.document.add_paragraph()

        # Team members
        p = self.document.add_paragraph()
        run = p.add_run("Team Members:")
        run.font.bold = True

        students = content.get("students", [])
        if students:
            for idx, student in enumerate(students, 1):
                name = student.get("name", "Student") if isinstance(student, dict) else str(student)
                self.document.add_paragraph(f"{idx}. {name}")

        self.document.add_paragraph()

        # Date and Place
        date_str = content.get("date", datetime.now().strftime("%B %Y"))
        p = self.document.add_paragraph()
        p.add_run(f"Date: {date_str}")

        p = self.document.add_paragraph()
        p.add_run(f"Place: {college_name}")

        self.document.add_page_break()

    def _add_auto_section(self, section_id: str):
        """Add auto-generated section (TOC, list of figures, etc.)"""

        if section_id == "toc":
            self._add_toc()
        elif section_id == "list_figures":
            self._add_list_of_figures()
        elif section_id == "list_tables":
            self._add_list_of_tables()

    def _add_toc(self):
        """Add Table of Contents with auto-update field"""
        p = self.document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("TABLE OF CONTENTS")
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.name = 'Arial'

        self.document.add_paragraph()

        # Add instruction note
        note_p = self.document.add_paragraph()
        note_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        note_run = note_p.add_run("(Press Ctrl+A then F9 to update all fields, or right-click TOC and select 'Update Field')")
        note_run.font.italic = True
        note_run.font.size = Pt(10)
        note_run.font.color.rgb = RGBColor(128, 128, 128)

        self.document.add_paragraph()

        # Add TOC field with proper Word field codes
        paragraph = self.document.add_paragraph()
        run = paragraph.add_run()

        # Begin field
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        run._r.append(fldChar1)

        # Field instruction - TOC with heading levels 1-3, hyperlinks, and page numbers
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = ' TOC \\o "1-3" \\h \\z \\u '
        run._r.append(instrText)

        # Separate field
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        run._r.append(fldChar2)

        # Placeholder content (will be replaced when TOC is updated)
        placeholder_entries = [
            ("CHAPTER 1: INTRODUCTION", "1"),
            ("CHAPTER 2: LITERATURE REVIEW", "9"),
            ("CHAPTER 3: REQUIREMENT ANALYSIS", "19"),
            ("CHAPTER 4: SYSTEM DESIGN", "31"),
            ("CHAPTER 5: IMPLEMENTATION", "43"),
            ("CHAPTER 6: TESTING", "53"),
            ("CHAPTER 7: CONCLUSION", "61"),
            ("REFERENCES", "65"),
            ("APPENDICES", "68"),
        ]

        for entry, page in placeholder_entries:
            entry_p = self.document.add_paragraph()
            entry_p.paragraph_format.tab_stops.add_tab_stop(Inches(6), alignment=WD_ALIGN_PARAGRAPH.RIGHT, leader=1)  # Dot leader
            entry_run = entry_p.add_run(f"{entry}\t{page}")
            entry_run.font.size = Pt(12)
            entry_run.font.name = 'Times New Roman'

        # End field
        end_p = self.document.add_paragraph()
        end_run = end_p.add_run()
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')
        end_run._r.append(fldChar3)

        self.document.add_page_break()

    def _add_list_of_figures(self):
        """Add List of Figures with auto-update field"""
        p = self.document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("LIST OF FIGURES")
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.name = 'Arial'

        self.document.add_paragraph()

        # Add instruction note
        note_p = self.document.add_paragraph()
        note_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        note_run = note_p.add_run("(Press Ctrl+A then F9 to update, or right-click and select 'Update Field')")
        note_run.font.italic = True
        note_run.font.size = Pt(10)
        note_run.font.color.rgb = RGBColor(128, 128, 128)

        self.document.add_paragraph()

        # Add Table of Figures field
        paragraph = self.document.add_paragraph()
        run = paragraph.add_run()

        # Begin field
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        run._r.append(fldChar1)

        # Field instruction for Table of Figures
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = ' TOC \\h \\z \\c "Figure" '
        run._r.append(instrText)

        # Separate field
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        run._r.append(fldChar2)

        # Placeholder entries for figures
        figure_entries = [
            ("Figure 4.1: Use Case Diagram", "33"),
            ("Figure 4.2: Class Diagram", "35"),
            ("Figure 4.3: Sequence Diagram", "36"),
            ("Figure 4.4: Activity Diagram", "38"),
            ("Figure 4.5: ER Diagram", "40"),
            ("Figure 5.1: System Architecture", "44"),
            ("Figure 6.1: Login Screen", "54"),
            ("Figure 6.2: Dashboard Screen", "55"),
        ]

        for entry, page in figure_entries:
            entry_p = self.document.add_paragraph()
            entry_p.paragraph_format.tab_stops.add_tab_stop(Inches(6), alignment=WD_ALIGN_PARAGRAPH.RIGHT, leader=1)
            entry_run = entry_p.add_run(f"{entry}\t{page}")
            entry_run.font.size = Pt(12)
            entry_run.font.name = 'Times New Roman'

        # End field
        end_p = self.document.add_paragraph()
        end_run = end_p.add_run()
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')
        end_run._r.append(fldChar3)

        self.document.add_page_break()

    def _add_list_of_tables(self):
        """Add List of Tables with auto-update field"""
        p = self.document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("LIST OF TABLES")
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.name = 'Arial'

        self.document.add_paragraph()

        # Add instruction note
        note_p = self.document.add_paragraph()
        note_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        note_run = note_p.add_run("(Press Ctrl+A then F9 to update, or right-click and select 'Update Field')")
        note_run.font.italic = True
        note_run.font.size = Pt(10)
        note_run.font.color.rgb = RGBColor(128, 128, 128)

        self.document.add_paragraph()

        # Add Table of Tables field
        paragraph = self.document.add_paragraph()
        run = paragraph.add_run()

        # Begin field
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        run._r.append(fldChar1)

        # Field instruction for Table of Tables
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = ' TOC \\h \\z \\c "Table" '
        run._r.append(instrText)

        # Separate field
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        run._r.append(fldChar2)

        # Placeholder entries for tables
        table_entries = [
            ("Table 2.1: Comparison of Existing Systems", "15"),
            ("Table 3.1: Functional Requirements", "25"),
            ("Table 3.2: Non-Functional Requirements", "28"),
            ("Table 4.1: Database Schema - Users", "41"),
            ("Table 4.2: Database Schema - Projects", "42"),
            ("Table 5.1: Hardware Requirements", "43"),
            ("Table 5.2: Software Requirements", "44"),
            ("Table 6.1: Test Cases Summary", "56"),
        ]

        for entry, page in table_entries:
            entry_p = self.document.add_paragraph()
            entry_p.paragraph_format.tab_stops.add_tab_stop(Inches(6), alignment=WD_ALIGN_PARAGRAPH.RIGHT, leader=1)
            entry_run = entry_p.add_run(f"{entry}\t{page}")
            entry_run.font.size = Pt(12)
            entry_run.font.name = 'Times New Roman'

        # End field
        end_p = self.document.add_paragraph()
        end_run = end_p.add_run()
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')
        end_run._r.append(fldChar3)

        self.document.add_page_break()

    def _add_generated_section(self, section_id: str, title: str, content: Dict, project_data: Dict = None):
        """Add AI-generated section content"""

        # Handle case where content might be a string instead of dict
        if isinstance(content, str):
            content = {"content": content}
        elif not isinstance(content, dict):
            content = {}

        # Chapter/Section heading
        if section_id.startswith("ch"):
            # Chapter heading
            p = self.document.add_paragraph(title, style='Heading 1')
            p.paragraph_format.page_break_before = True
        else:
            # Section heading
            p = self.document.add_paragraph(title, style='Heading 1')

        # Add main content
        main_content = content.get("content", "")

        # If no main content but has error/fallback flag, generate placeholder
        if not main_content and content.get("fallback"):
            main_content = f"""This section covers {title}. The {title.lower()} provides essential information about the project implementation, covering key aspects of the system design and development process.

The content in this section is organized to give readers a comprehensive understanding of the technical aspects involved. Each subsection addresses specific components and their roles in the overall system architecture.

Key points covered include the design decisions, implementation strategies, and best practices followed during development. The technical details are presented in a manner that facilitates understanding for both technical and non-technical stakeholders."""
            logger.warning(f"[WordGenerator] Using fallback content for section: {section_id}")

        if main_content:
            self._add_formatted_text(main_content)

        # Add subsections with integrated diagrams
        subsections = content.get("subsections", [])
        for subsection in subsections:
            if isinstance(subsection, dict):
                sub_title = subsection.get("title", subsection.get("id", "Subsection"))
                sub_content = subsection.get("content", "")
            else:
                sub_title = str(subsection)
                sub_content = ""

            # Subsection heading
            if sub_title:
                p = self.document.add_paragraph(sub_title, style='Heading 2')

            # Subsection content
            if sub_content:
                self._add_formatted_text(sub_content)

            # Add diagram if this subsection is about a specific diagram type
            if project_data:
                self._add_diagram_for_subsection(sub_title, section_id, project_data)

        # If no content was added at all, add a basic placeholder
        if not main_content and not subsections:
            logger.warning(f"[WordGenerator] No content for section {section_id}, adding placeholder")
            placeholder_text = f"""This section presents the details of {title.replace('Chapter', '').strip()}. The information provided here is crucial for understanding the project's development and implementation.

The section includes relevant technical details, design decisions, and implementation strategies that were employed during the development process. Each aspect is carefully documented to ensure clarity and completeness."""
            self._add_formatted_text(placeholder_text)

        # Add figures
        figures = content.get("figures", [])
        for figure in figures:
            self._add_figure_placeholder(figure)

        # Add tables
        tables = content.get("tables", [])
        for table in tables:
            self._add_table(table)

        # Handle Q&A pairs for viva documents
        qa_pairs = content.get("qa_pairs", [])
        if qa_pairs:
            for idx, qa in enumerate(qa_pairs, 1):
                if isinstance(qa, dict):
                    question = qa.get("question", "")
                    answer = qa.get("answer", "")
                    follow_up = qa.get("follow_up_tips", "")

                    if question:
                        # Question heading
                        q_para = self.document.add_paragraph()
                        q_para.paragraph_format.space_before = Pt(12)
                        q_para.paragraph_format.space_after = Pt(6)
                        q_run = q_para.add_run(f"Q{idx}: {question}")
                        q_run.bold = True
                        q_run.font.size = Pt(12)
                        q_run.font.name = 'Times New Roman'
                        q_run.font.color.rgb = RGBColor(0, 51, 102)  # Dark blue for questions

                    if answer:
                        # Answer with prefix
                        a_para = self.document.add_paragraph()
                        a_para.paragraph_format.left_indent = Cm(0.5)
                        a_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
                        a_para.paragraph_format.space_after = Pt(6)
                        a_prefix = a_para.add_run("Answer: ")
                        a_prefix.bold = True
                        a_prefix.font.size = Pt(12)
                        a_prefix.font.name = 'Times New Roman'
                        self._add_run_with_formatting(a_para, answer)
                        for run in a_para.runs[1:]:  # Skip the "Answer:" prefix
                            run.font.size = Pt(12)
                            run.font.name = 'Times New Roman'

                    if follow_up:
                        # Follow-up tips in italics
                        tip_para = self.document.add_paragraph()
                        tip_para.paragraph_format.left_indent = Cm(0.5)
                        tip_para.paragraph_format.space_after = Pt(12)
                        tip_prefix = tip_para.add_run("Follow-up Tips: ")
                        tip_prefix.bold = True
                        tip_prefix.italic = True
                        tip_prefix.font.size = Pt(11)
                        tip_prefix.font.name = 'Times New Roman'
                        tip_content = tip_para.add_run(follow_up)
                        tip_content.italic = True
                        tip_content.font.size = Pt(11)
                        tip_content.font.name = 'Times New Roman'
                        tip_content.font.color.rgb = RGBColor(80, 80, 80)  # Gray color

    def _add_formatted_text(self, text: str):
        """Add text with markdown-like formatting - supports bullet points, numbered lists, and paragraphs"""
        # Split into paragraphs (double newline or single newline for list items)
        paragraphs = text.split('\n\n')

        for para_text in paragraphs:
            if not para_text.strip():
                continue

            # Handle multi-line content within a paragraph block
            lines = para_text.strip().split('\n')

            in_list = False
            for line in lines:
                line_stripped = line.strip()
                if not line_stripped:
                    continue

                # Check for bullet points (-, •, *, or bullet unicode)
                if line_stripped.startswith(('-', '•', '*', '●', '○', '▪', '→')):
                    in_list = True
                    # Extract bullet content
                    bullet_content = line_stripped.lstrip('-•*●○▪→ ').strip()
                    if bullet_content:
                        p = self.document.add_paragraph(style='List Bullet')
                        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
                        p.paragraph_format.space_after = Pt(6)
                        self._add_run_with_formatting(p, bullet_content)
                        # Ensure 12pt font for bullet text
                        for run in p.runs:
                            run.font.size = Pt(12)
                            run.font.name = 'Times New Roman'

                # Check for numbered lists (1. 2. 3. etc)
                elif re.match(r'^\d+[\.\)]\s', line_stripped):
                    in_list = True
                    # Extract number and content
                    match = re.match(r'^(\d+)[\.\)]\s*(.+)', line_stripped)
                    if match:
                        num_content = match.group(2).strip()
                        p = self.document.add_paragraph(style='List Number')
                        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
                        p.paragraph_format.space_after = Pt(6)
                        self._add_run_with_formatting(p, num_content)
                        # Ensure 12pt font
                        for run in p.runs:
                            run.font.size = Pt(12)
                            run.font.name = 'Times New Roman'

                # Regular paragraph text
                else:
                    p = self.document.add_paragraph()
                    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
                    p.paragraph_format.space_after = Pt(8)
                    if not in_list:
                        p.paragraph_format.first_line_indent = Cm(1.27)

                    # Handle inline formatting
                    self._add_run_with_formatting(p, line_stripped)

                    # Ensure 12pt font for body text
                    for run in p.runs:
                        run.font.size = Pt(12)
                        run.font.name = 'Times New Roman'

                    in_list = False

    def _add_run_with_formatting(self, paragraph, text: str):
        """Add text run with bold/italic/code formatting"""
        # Simple markdown-like parsing
        parts = re.split(r'(\*\*.*?\*\*|`.*?`)', text)

        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            elif part.startswith('`') and part.endswith('`'):
                run = paragraph.add_run(part[1:-1])
                run.font.name = 'Consolas'
                run.font.size = Pt(10)
            else:
                paragraph.add_run(part)

    def _add_figure_placeholder(self, figure: Dict):
        """Add figure placeholder with explanatory text"""
        description = figure.get('description', 'Image')
        caption = figure.get('caption', 'Figure caption')

        # Diagram explanations for common diagram types
        diagram_explanations = {
            "use case": "This Use Case Diagram illustrates the interactions between actors (users) and the system. It shows the main functionalities available to different user types and how they interact with the system to achieve their goals.",
            "class diagram": "This Class Diagram represents the static structure of the system, showing classes, their attributes, methods, and relationships. It provides a blueprint for implementing the object-oriented design of the application.",
            "sequence diagram": "This Sequence Diagram depicts the order of interactions between objects over time. It shows how different components communicate to complete a specific use case or process flow.",
            "activity diagram": "This Activity Diagram models the workflow of the system, showing the sequence of activities and decision points. It helps visualize the business logic and process flows within the application.",
            "er diagram": "This Entity-Relationship Diagram shows the database structure, including entities (tables), their attributes, and relationships. It provides a visual representation of how data is organized and connected.",
            "entity-relationship": "This Entity-Relationship Diagram shows the database structure, including entities (tables), their attributes, and relationships. It provides a visual representation of how data is organized and connected.",
            "dfd": "This Data Flow Diagram illustrates how data moves through the system, showing processes, data stores, and external entities. It helps understand the flow of information within the application.",
            "data flow": "This Data Flow Diagram illustrates how data moves through the system, showing processes, data stores, and external entities. It helps understand the flow of information within the application.",
            "architecture": "This Architecture Diagram shows the high-level structure of the system, including components, layers, and their interactions. It provides an overview of how the system is organized and deployed.",
            "system architecture": "This System Architecture Diagram shows the overall structure including frontend, backend, database, and external services. It illustrates the technology stack and how components communicate.",
        }

        # Find matching explanation
        explanation = None
        desc_lower = description.lower()
        for key, expl in diagram_explanations.items():
            if key in desc_lower:
                explanation = expl
                break

        # Default explanation if no match found
        if not explanation:
            explanation = f"This diagram provides a visual representation of {description.lower()}. It helps in understanding the structure and relationships within the system component being illustrated."

        # Add explanation text before the placeholder
        exp_p = self.document.add_paragraph()
        exp_p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        exp_p.paragraph_format.space_after = Pt(8)
        exp_run = exp_p.add_run(explanation)
        exp_run.font.size = Pt(12)
        exp_run.font.name = 'Times New Roman'

        self.document.add_paragraph()

        # Add placeholder box with border styling
        p = self.document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Create a visual placeholder box
        placeholder_text = f"┌{'─' * 50}┐\n│{' ' * 20}DIAGRAM PLACEHOLDER{' ' * 11}│\n│{' ' * 50}│\n│{' ' * 10}{description[:30]:^30}{' ' * 10}│\n│{' ' * 50}│\n│{' ' * 5}(Insert actual diagram image here){' ' * 10}│\n└{'─' * 50}┘"
        run = p.add_run(placeholder_text)
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(100, 100, 100)

        # Add caption with proper formatting
        cap_p = self.document.add_paragraph()
        cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_run = cap_p.add_run(f"Figure: {caption}")
        cap_run.font.italic = True
        cap_run.font.size = Pt(10)
        cap_run.font.name = 'Times New Roman'

        self.document.add_paragraph()

    def _add_table(self, table_data: Dict):
        """Add table to document"""
        headers = table_data.get("headers", [])
        rows = table_data.get("rows", [])
        caption = table_data.get("caption", "Table")

        if not headers or not rows:
            return

        # Create table
        table = self.document.add_table(rows=len(rows) + 1, cols=len(headers))
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Header row
        header_row = table.rows[0]
        for i, header in enumerate(headers):
            cell = header_row.cells[i]
            cell.text = str(header)
            # Bold header
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        # Data rows
        for row_idx, row_data in enumerate(rows):
            row = table.rows[row_idx + 1]
            for col_idx, cell_data in enumerate(row_data):
                if col_idx < len(row.cells):
                    row.cells[col_idx].text = str(cell_data)

        # Add caption
        p = self.document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(f"Table: {caption}")
        p.runs[0].font.italic = True
        p.runs[0].font.size = Pt(10)

        self.document.add_paragraph()

    def _add_code_section(self, content: Dict):
        """Add code appendix section"""
        p = self.document.add_paragraph("APPENDIX: SOURCE CODE", style='Heading 1')
        p.paragraph_format.page_break_before = True

        files = content.get("files", [])

        for file in files:
            filename = file.get("filename", "code.txt")
            language = file.get("language", "")
            code = file.get("content", "")

            # File header
            p = self.document.add_paragraph()
            run = p.add_run(f"File: {filename}")
            run.bold = True
            run.font.size = Pt(11)

            # Code block
            p = self.document.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.5)
            run = p.add_run(code[:3000])  # Limit code length
            run.font.name = 'Consolas'
            run.font.size = Pt(9)

            self.document.add_paragraph()

    def _update_toc_fields(self):
        """Mark document to update TOC fields when opened"""
        # This tells Word to update fields when document is opened
        settings_element = self.document.settings.element
        update_fields = OxmlElement('w:updateFields')
        update_fields.set(qn('w:val'), 'true')
        settings_element.append(update_fields)

    def _add_diagram_for_subsection(self, sub_title: str, section_id: str, project_data: Dict):
        """Add diagram inline within a subsection based on the subsection title"""
        sub_title_lower = sub_title.lower()

        # Mapping of subsection titles to diagram types
        # Only add diagrams to specific subsections, NOT to UI/UX Design
        diagram_mapping = {
            # Chapter 3: Requirements
            "use case": ("use_case", "Use Case Diagram"),
            "data flow": ("dfd_0", "Data Flow Diagram (Level 0)"),
            "dfd": ("dfd_0", "Data Flow Diagram (Level 0)"),

            # Chapter 4: Design - Only for specific diagram sections
            "er diagram": ("er", "Entity-Relationship Diagram"),
            "entity-relationship": ("er", "Entity-Relationship Diagram"),
            "class diagram": ("class", "Class Diagram"),
            "sequence diagram": ("sequence", "Sequence Diagram"),
            "activity diagram": ("activity", "Activity Diagram"),
            "system architecture": ("system_architecture", "System Architecture Diagram"),
            "architecture diagram": ("system_architecture", "System Architecture Diagram"),
        }

        # Skip UI/UX Design - it should NOT have UML diagrams
        if "ui" in sub_title_lower or "ux" in sub_title_lower or "user interface" in sub_title_lower:
            return

        # Skip API Design - it should NOT have UML diagrams
        if "api" in sub_title_lower:
            return

        # Skip Database Design - text descriptions only (ER diagram is separate)
        if "database design" in sub_title_lower and "er" not in sub_title_lower:
            return

        # Find matching diagram
        for keyword, (diagram_type, diagram_title) in diagram_mapping.items():
            if keyword in sub_title_lower:
                self._add_inline_diagram(diagram_type, diagram_title, project_data)
                break

    def _add_inline_diagram(self, diagram_type: str, diagram_title: str, project_data: Dict):
        """Add a diagram inline within the current section"""
        try:
            diagram_path = self._generate_diagram(diagram_type, project_data)

            # Check if it's a real file path (not a placeholder string)
            if diagram_path and not diagram_path.startswith('[') and os.path.exists(diagram_path):
                # Add the diagram image
                p = self.document.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(diagram_path, width=Inches(5.5))

                # Add caption using proper SEQ field for figure numbering
                cap_p = self.document.add_paragraph()
                cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Add SEQ field for auto-numbering
                cap_run = cap_p.add_run()
                fldChar1 = OxmlElement('w:fldChar')
                fldChar1.set(qn('w:fldCharType'), 'begin')
                cap_run._r.append(fldChar1)

                instrText = OxmlElement('w:instrText')
                instrText.text = ' SEQ Figure \\* ARABIC '
                cap_run._r.append(instrText)

                fldChar2 = OxmlElement('w:fldChar')
                fldChar2.set(qn('w:fldCharType'), 'separate')
                cap_run._r.append(fldChar2)

                num_run = cap_p.add_run("X")  # Placeholder number
                num_run.font.italic = True
                num_run.font.size = Pt(10)

                fldChar3 = OxmlElement('w:fldChar')
                fldChar3.set(qn('w:fldCharType'), 'end')
                cap_run._r.append(fldChar3)

                # Add caption text
                text_run = cap_p.add_run(f": {diagram_title}")
                text_run.font.italic = True
                text_run.font.size = Pt(10)
                text_run.font.name = 'Times New Roman'

                self.document.add_paragraph()  # Add spacing
                logger.info(f"[WordGenerator] Added inline {diagram_title}")
            else:
                # Add placeholder with detailed explanation if generation failed
                logger.warning(f"[WordGenerator] Diagram not generated for {diagram_title}, adding placeholder")
                self._add_figure_placeholder({
                    "description": diagram_title,
                    "caption": diagram_title
                })
        except Exception as e:
            logger.error(f"[WordGenerator] Error adding inline diagram {diagram_title}: {e}")
            # Still add placeholder on error
            self._add_figure_placeholder({
                "description": diagram_title,
                "caption": diagram_title
            })

    def _add_diagram_section(self, title: str, project_data: Dict, diagram_type: str):
        """Add a specific UML diagram to the document"""
        try:
            # Add subsection heading
            p = self.document.add_paragraph(title, style='Heading 3')

            # Generate the diagram
            diagram_path = self._generate_diagram(diagram_type, project_data)

            if diagram_path and os.path.exists(diagram_path):
                # Add the diagram image
                p = self.document.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(diagram_path, width=Inches(5.5))

                # Add caption
                p = self.document.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(f"Figure: {title}")
                run.font.italic = True
                run.font.size = Pt(10)

                self.document.add_paragraph()  # Add spacing

                logger.info(f"[WordGenerator] Added {title} diagram")
            else:
                # Add placeholder if diagram generation failed
                self._add_figure_placeholder({
                    "description": title,
                    "caption": title
                })

        except Exception as e:
            logger.error(f"[WordGenerator] Error adding diagram {title}: {e}")
            self._add_figure_placeholder({
                "description": title,
                "caption": f"{title} (generation failed)"
            })

    def _generate_diagram(self, diagram_type: str, project_data: Dict) -> Optional[str]:
        """Generate a specific diagram using UML generator"""
        try:
            # Check if diagram was pre-generated by chunked_document_agent
            pre_generated = project_data.get('generated_diagrams', {})
            if diagram_type in pre_generated:
                return pre_generated[diagram_type]

            project_name = project_data.get('project_name', 'System')
            # Use stored project_id and user_id for isolation
            project_id = getattr(self, 'project_id', None) or project_data.get('project_id')
            user_id = getattr(self, 'user_id', None) or project_data.get('user_id')
            features = project_data.get('features', [])
            database_tables = project_data.get('database_tables', [])

            if diagram_type == "use_case":
                # Generate use case diagram
                actors = ['User', 'Admin']
                if 'authentication' in str(features).lower():
                    actors.append('Guest')
                use_cases = features[:8] if features else ['Login', 'Register', 'View Dashboard', 'Manage Data']
                return uml_generator.generate_use_case_diagram(
                    project_name=project_name,
                    actors=actors,
                    use_cases=use_cases,
                    project_id=project_id,
                    user_id=user_id
                )

            elif diagram_type == "class":
                # Generate class diagram
                classes = self._extract_classes(project_data)
                return uml_generator.generate_class_diagram(classes, project_id=project_id, user_id=user_id)

            elif diagram_type == "sequence":
                # Generate sequence diagram
                participants = ['User', 'Frontend', 'API', 'Database']
                messages = [
                    {'from': 'User', 'to': 'Frontend', 'message': 'Submit Request'},
                    {'from': 'Frontend', 'to': 'API', 'message': 'API Call'},
                    {'from': 'API', 'to': 'Database', 'message': 'Query'},
                    {'from': 'Database', 'to': 'API', 'message': 'Result', 'type': 'return'},
                    {'from': 'API', 'to': 'Frontend', 'message': 'Response', 'type': 'return'},
                    {'from': 'Frontend', 'to': 'User', 'message': 'Display', 'type': 'return'},
                ]
                return uml_generator.generate_sequence_diagram(participants, messages, project_id=project_id, user_id=user_id)

            elif diagram_type == "activity":
                # Generate activity diagram
                activities = [
                    'Start Application',
                    'User Authentication',
                    'Load Dashboard',
                    'Process User Request',
                    'Update Database',
                    'Return Response'
                ]
                return uml_generator.generate_activity_diagram(activities, project_id=project_id, user_id=user_id)

            elif diagram_type == "er":
                # Generate ER diagram
                entities = self._extract_entities(project_data)
                return uml_generator.generate_er_diagram(entities, project_id=project_id, user_id=user_id)

            elif diagram_type == "dfd_0":
                # Generate DFD Level 0
                return uml_generator.generate_dfd(
                    level=0,
                    processes=[project_name],
                    data_stores=['Database'],
                    external_entities=['User', 'Admin'],
                    data_flows=[
                        {'from': 'User', 'to': project_name, 'data': 'Request'},
                        {'from': project_name, 'to': 'User', 'data': 'Response'},
                        {'from': project_name, 'to': 'Database', 'data': 'Query'},
                    ],
                    project_id=project_id,
                    user_id=user_id
                )

            elif diagram_type == "system_architecture":
                # Generate System Architecture Diagram
                return uml_generator.generate_system_architecture_diagram(project_data, project_id=project_id, user_id=user_id)

            return None

        except Exception as e:
            logger.error(f"[WordGenerator] Error generating {diagram_type} diagram: {e}")
            return None

    def _extract_classes(self, project_data: Dict) -> List[Dict]:
        """Extract class information from project data for class diagram"""
        classes = []

        # From database tables
        tables = project_data.get('database_tables', [])
        for table in tables[:5]:
            classes.append({
                'name': table.title() if isinstance(table, str) else table.get('name', 'Entity'),
                'attributes': ['id', 'name', 'created_at', 'updated_at'],
                'methods': ['create', 'read', 'update', 'delete'],
                'relationships': []
            })

        # Default classes if none found
        if not classes:
            classes = [
                {
                    'name': 'User',
                    'attributes': ['id', 'name', 'email', 'password'],
                    'methods': ['login', 'logout', 'register'],
                    'relationships': []
                },
                {
                    'name': 'Controller',
                    'attributes': ['routes', 'middleware'],
                    'methods': ['handleRequest', 'validateInput'],
                    'relationships': [{'target': 'Service', 'type': 'association'}]
                },
                {
                    'name': 'Service',
                    'attributes': ['repository'],
                    'methods': ['processData', 'validateBusiness'],
                    'relationships': [{'target': 'Repository', 'type': 'association'}]
                },
                {
                    'name': 'Repository',
                    'attributes': ['database'],
                    'methods': ['find', 'save', 'delete'],
                    'relationships': []
                }
            ]

        return classes

    def _extract_entities(self, project_data: Dict) -> List[Dict]:
        """Extract entity information for ER diagram"""
        entities = []

        tables = project_data.get('database_tables', [])
        for table in tables[:6]:
            name = table if isinstance(table, str) else table.get('name', 'Entity')
            entities.append({
                'name': name,
                'attributes': ['id', 'name', 'description', 'created_at', 'updated_at'],
                'primary_key': 'id'
            })

        # Default entities
        if not entities:
            entities = [
                {'name': 'User', 'attributes': ['id', 'name', 'email', 'password_hash'], 'primary_key': 'id'},
                {'name': 'Project', 'attributes': ['id', 'title', 'description', 'user_id'], 'primary_key': 'id'},
                {'name': 'Document', 'attributes': ['id', 'name', 'content', 'project_id'], 'primary_key': 'id'},
            ]

        return entities
